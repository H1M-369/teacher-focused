import os
import json
import streamlit as st
from io import BytesIO
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pdf_utils import extract_text_from_pdf
from creator import generate_assignment

load_dotenv()

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(page_title="AI Assignment Creator", page_icon="📋", layout="wide")
st.title("📋 AI Assignment Creator")
st.caption("Enter topics and total marks — AI builds a balanced assignment and allocates marks.")

# ── Session state ─────────────────────────────────────────────────────────────
for key, default in {
    "assignment": None,
    "generated": False,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input(
        "Claude API Key",
        type="password",
        value=os.environ.get("ANTHROPIC_API_KEY", ""),
        help="Get yours free at console.anthropic.com",
    )
    if api_key:
        os.environ["ANTHROPIC_API_KEY"] = api_key
        col_a, col_b = st.columns(2)
        col_a.success("Key ready")
        if col_b.button("💾 Save", help="Save to .env so you don't retype it"):
            Path(".env").write_text(f"ANTHROPIC_API_KEY={api_key}\n", encoding="utf-8")
            st.success("Saved")
    else:
        st.warning("Enter your API key to continue")
        st.markdown("[Get a free API key →](https://console.anthropic.com)")

    st.divider()
    st.markdown("**How marks are allocated**")
    st.caption(
        "• Topics are ranked by complexity from your reference material.\n"
        "• Broader/harder topics get more marks and more questions.\n"
        "• You can lock marks to specific topics before generating.\n"
        "• Total always equals exactly what you set."
    )


def require_api_key():
    if not os.environ.get("ANTHROPIC_API_KEY"):
        st.error("Enter your Claude API key in the sidebar first.")
        st.stop()


# ── Word document builder ─────────────────────────────────────────────────────
def build_word_doc(data: dict, course: str, time_allowed: str, instructions: str) -> bytes:
    doc = Document()

    # Title block
    title_para = doc.add_heading(data.get("assignment_title", "Assignment"), 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if course:
        c = doc.add_paragraph(course)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].bold = True

    doc.add_paragraph()

    # Meta info table
    meta = doc.add_table(rows=1, cols=2)
    meta.style = "Table Grid"
    row = meta.rows[0].cells
    left_text = f"Total Marks: {data['total_marks']}"
    right_text = f"Time Allowed: {time_allowed}" if time_allowed else ""
    row[0].text = left_text
    row[1].text = right_text
    row[0].paragraphs[0].runs[0].bold = True
    if right_text:
        row[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    if instructions:
        doc.add_heading("Instructions", level=2)
        doc.add_paragraph(instructions)
        doc.add_paragraph()

    # Questions
    for topic_data in data["topics"]:
        heading = doc.add_heading(
            f"{topic_data['topic']}  [{topic_data['marks_allocated']} marks]", level=2
        )

        for q in topic_data["questions"]:
            p = doc.add_paragraph(style="List Number")
            run_q = p.add_run(q["question"])
            run_q.font.size = Pt(11)
            run_m = p.add_run(f"  ({q['marks']} marks)")
            run_m.bold = True
            run_m.font.size = Pt(11)

        doc.add_paragraph()  # spacing between topics

    # Rationale at end (for teacher reference)
    doc.add_page_break()
    doc.add_heading("Mark Allocation Rationale (Teacher Copy)", level=2)
    doc.add_paragraph(data.get("mark_allocation_rationale", ""))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# SETUP FORM
# ═══════════════════════════════════════════════════════════════════════════════
st.header("Assignment Details")

col1, col2 = st.columns(2)
with col1:
    assignment_title = st.text_input("Assignment Title (optional — AI will suggest one if blank)")
    course_name = st.text_input("Course / Subject Name")
with col2:
    total_marks = st.number_input("Total Marks", min_value=10, max_value=500, value=100, step=5)
    time_allowed = st.text_input("Time Allowed (optional)", placeholder="e.g. 2 hours")

instructions = st.text_area(
    "Student Instructions (optional)",
    placeholder="e.g. Answer ALL questions. Show all working where applicable.",
    height=80,
)

st.divider()

# ── Reference material ────────────────────────────────────────────────────────
st.subheader("Reference / Unit Material (optional but recommended)")
st.caption("Upload your textbook or lecture notes — questions will be grounded in your course content.")
reference_files = st.file_uploader(
    "Upload PDFs", type="pdf", accept_multiple_files=True, key="refs"
)

st.divider()

# ── Topics ────────────────────────────────────────────────────────────────────
st.subheader("Topics")
st.caption(
    "Add all topics the assignment should cover. "
    "Optionally lock marks to a topic — otherwise AI distributes them by complexity."
)

# Dynamic topic list stored in session state
if "topics" not in st.session_state:
    st.session_state.topics = [{"topic": "", "locked_marks": None}]


def add_topic():
    st.session_state.topics.append({"topic": "", "locked_marks": None})


def remove_topic(i):
    st.session_state.topics.pop(i)


for i, t in enumerate(st.session_state.topics):
    col_t, col_m, col_del = st.columns([5, 2, 1])
    with col_t:
        st.session_state.topics[i]["topic"] = st.text_input(
            f"Topic {i + 1}",
            value=t["topic"],
            key=f"topic_{i}",
            label_visibility="collapsed",
            placeholder=f"Topic {i + 1} name…",
        )
    with col_m:
        lock = st.number_input(
            "Lock marks",
            min_value=0,
            max_value=int(total_marks),
            value=int(t["locked_marks"]) if t["locked_marks"] else 0,
            key=f"lock_{i}",
            label_visibility="collapsed",
            help="Set to 0 to let AI decide",
        )
        st.session_state.topics[i]["locked_marks"] = lock if lock > 0 else None
    with col_del:
        if len(st.session_state.topics) > 1:
            if st.button("✕", key=f"del_{i}", help="Remove topic"):
                remove_topic(i)
                st.rerun()

st.button("＋ Add Topic", on_click=add_topic)

# Locked marks validation
locked_sum = sum(t["locked_marks"] or 0 for t in st.session_state.topics)
if locked_sum > total_marks:
    st.error(f"Locked marks total ({locked_sum}) exceeds total marks ({total_marks}). Reduce some locked values.")
elif locked_sum > 0:
    st.info(f"Locked marks: {locked_sum} / {total_marks}. AI will distribute the remaining {total_marks - locked_sum} marks.")

st.divider()

# ═══════════════════════════════════════════════════════════════════════════════
# GENERATE
# ═══════════════════════════════════════════════════════════════════════════════
valid_topics = [t for t in st.session_state.topics if t["topic"].strip()]
can_generate = (
    len(valid_topics) >= 1
    and locked_sum <= total_marks
    and bool(os.environ.get("ANTHROPIC_API_KEY"))
)

if st.button(
    "⚡ Generate Assignment",
    type="primary",
    use_container_width=True,
    disabled=not can_generate,
):
    require_api_key()
    with st.spinner("Analysing topics, allocating marks, writing questions… this takes 30–60 seconds."):
        try:
            ref_texts = [extract_text_from_pdf(f) for f in reference_files] if reference_files else []
            data = generate_assignment(
                topics=valid_topics,
                total_marks=int(total_marks),
                reference_texts=ref_texts,
                title=assignment_title,
                course=course_name,
            )
            st.session_state.assignment = data
            st.session_state.generated = True
            st.rerun()
        except Exception as e:
            st.error(f"Error: {e}")

if not can_generate and not os.environ.get("ANTHROPIC_API_KEY"):
    st.caption("Enter your API key in the sidebar to enable generation.")
elif not can_generate:
    st.caption("Add at least one topic to generate.")


# ═══════════════════════════════════════════════════════════════════════════════
# PREVIEW & EXPORT
# ═══════════════════════════════════════════════════════════════════════════════
if st.session_state.assignment:
    data = st.session_state.assignment

    st.divider()
    st.header("Generated Assignment")

    # Summary metrics
    q_count = sum(len(t["questions"]) for t in data["topics"])
    c1, c2, c3 = st.columns(3)
    c1.metric("Total Questions", q_count)
    c2.metric("Total Marks", data["total_marks"])
    c3.metric("Topics Covered", len(data["topics"]))

    # Mark allocation rationale
    with st.expander("📊 How marks were allocated"):
        st.write(data.get("mark_allocation_rationale", ""))
        st.divider()
        for t in data["topics"]:
            pct = round((t["marks_allocated"] / data["total_marks"]) * 100, 1)
            st.markdown(
                f"**{t['topic']}** — {t['marks_allocated']} marks ({pct}%)  \n"
                f"_{t.get('complexity_note', '')}_"
            )

    st.divider()

    # Editable preview
    st.subheader("Preview & Edit")
    st.caption("Edit questions or marks directly. Changes are reflected in the download.")

    updated_topics = []
    global_q_num = 1

    for ti, topic_data in enumerate(data["topics"]):
        with st.expander(
            f"{topic_data['topic']}  —  {topic_data['marks_allocated']} marks  "
            f"({len(topic_data['questions'])} questions)",
            expanded=True,
        ):
            updated_questions = []
            for qi, q in enumerate(topic_data["questions"]):
                col_q, col_m = st.columns([7, 1])
                with col_q:
                    new_q = st.text_area(
                        f"Q{global_q_num}",
                        value=q["question"],
                        key=f"q_{ti}_{qi}",
                        height=80,
                        label_visibility="collapsed",
                    )
                with col_m:
                    new_m = st.number_input(
                        "Marks",
                        min_value=1,
                        max_value=int(total_marks),
                        value=int(q["marks"]),
                        key=f"m_{ti}_{qi}",
                        label_visibility="collapsed",
                    )
                updated_questions.append({
                    **q,
                    "question_number": global_q_num,
                    "question": new_q,
                    "marks": new_m,
                })
                global_q_num += 1

            topic_total = sum(q["marks"] for q in updated_questions)
            updated_topics.append({
                **topic_data,
                "marks_allocated": topic_total,
                "questions": updated_questions,
            })

    # Live total check
    edited_total = sum(t["marks_allocated"] for t in updated_topics)
    if edited_total != total_marks:
        st.warning(f"Edited total: **{edited_total}** marks — target is **{int(total_marks)}**. Adjust question marks above.")
    else:
        st.success(f"Total marks: **{edited_total} / {int(total_marks)}** ✓")

    # Update session with edits
    data["topics"] = updated_topics
    data["total_marks"] = edited_total

    st.divider()

    # Downloads
    col_word, col_json = st.columns(2)

    with col_word:
        try:
            word_bytes = build_word_doc(data, course_name, time_allowed, instructions)
            safe_title = (assignment_title or data.get("assignment_title", "assignment")).replace(" ", "_")
            st.download_button(
                label="⬇️ Download Word Document (.docx)",
                data=word_bytes,
                file_name=f"{safe_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            st.caption("Formatted assignment ready to print or share with students.")
        except Exception as e:
            st.error(f"Word export error: {e}")

    with col_json:
        st.download_button(
            label="💾 Save Assignment (JSON)",
            data=json.dumps(data, indent=2),
            file_name=f"{safe_title}.json",
            mime="application/json",
            use_container_width=True,
        )
        st.caption("Save to reload or edit later without regenerating.")

    st.divider()
    if st.button("⚡ Regenerate (new questions)", use_container_width=True):
        st.session_state.assignment = None
        st.session_state.generated = False
        st.rerun()
